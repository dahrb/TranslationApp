from docx import Document
from deep_translator import (GoogleTranslator, DeeplTranslator)
import os
from docxcompose.composer import Composer

def read_in_docx():
    directory = '.'
    extension = '.docx'
    docs = []

    for a_file in os.listdir(directory):
        if a_file.endswith(extension):
            docs.append(a_file)

    try:
        docs = sorted(docs, key=lambda x: int("".join([i for i in x if i.isdigit()])))
    except:
        docs = sorted(docs)
        print('Warning: Because not all file names have numbers the sorting may be inaccurate. To correct please ensure all filenames have numbers')
        
    og = docs
    first_doc = docs[0]
    docs = docs[1:]

    return og, docs, first_doc

def combine_all_docx(filename_master,files_list,combined_name):
    number_of_sections=len(files_list)
    master = Document(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(files_list[i])
        composer.append(doc_temp)
    composer.save("{}.docx".format(combined_name))

# Function that enables the user to translate the entire word doc
def fullText(filename,lang, flag, flag2):
    
    if flag2:
        d = Document('{:}'.format(filename))
    else:
        d = Document('{:}.docx'.format(filename))
        
    fullText = []

    for i in range(0,len(d.paragraphs)):
    
        t = d.paragraphs[i].text
        
        try:
            if flag == '1':
                translation = GoogleTranslator(source='auto', target=lang).translate(text=t)
            else:
                translation = DeeplTranslator(api_key="1adeb428-2da3-52ec-0366-bfcf03ad27bb:fx", source="en", target=lang, use_free_api=True).translate(text=t)
            fullText.append(translation)
        except:
            pass

    fullerText = '\n\n'.join(fullText)
    
    return fullerText

#Function that enables the user to translate the doc inserting the 
def paraText(filename,lang):

    d = Document('{:}.docx'.format(filename))
    fullText = []
    
    for i in range(0,len(d.paragraphs)):
    
        fullText.append(d.paragraphs[i].text)
        t = d.paragraphs[i].text
        translation = GoogleTranslator(source='auto', target=lang).translate(text=t)
        fullText.append(translation)

    fullerText = '\n\n'.join(fullText)
    
    return fullerText

def menu():
    
    #Prints the menu
    print("------------------------------------------------------------------------------------------------------------------")
    print("\nDocument Translation Menu\n")
    print("1) Full Text Translation - Gives a new word document with a full translation of the text")
    print("2) Alternating Paragraph Translation - Gives a new word document alternating the source text with the translation")
    print("3) Combine - Merges multiple translation documents into one")

    print("q) Exit Program\n")
    
    #takes user input
    userSelection = input("Enter your choice followed by [ENTER]: ")
    print("------------------------------------------------------------------------------------------------------------------")

    return userSelection

#Main program
userSelection = menu()

while (userSelection != "q"):
    
    if (userSelection == "1"):
        
        print("What is the target language?")
        lang = input("\nEnter 'fr'- French, 'es' - Spanish, 'ca' - Catalan, 'en' - English, 'fi' - Finnish: - ")
    
        x = input("If you want to translate multiple documents type in 'M' otherwise proceed: ")
        
        if x == 'M': 
            
            userInput2 = input("\nPlease enter the prefix  you want the ouput doc to be saved as ")
            userInput3 = input("\nPlease enter the translation software you wish to use - 1: Google, 2: DeepL ")
            
            userInput,docs,first_doc = read_in_docx()
            
            counter = 1
            
            for i in userInput:
                p = fullText(i,lang,userInput3,True)
                doc = Document()
                doc.add_paragraph(p)
                doc.save('{:}{:}.docx'.format(userInput2,counter))
                counter += 1
                
        else:
            userInput = input("\nPlease enter the name of the source doc you want translating ")
            userInput3 = input("\nPlease enter the translation software you wish to use - 1: Google, 2: DeepL ")
            
            p = fullText(userInput,lang, userInput3, False)

            doc = Document()
            doc.add_paragraph(p)
        
            userInput2 = input("\nPlease enter the name you want the ouput doc to be saved as ")

            doc.save('{:}.docx'.format(userInput2))
            
        print("\nYour translated document can be found in the same folder as your source document")
        
    elif (userSelection == "2"):
        
        print("What is the target language?")
        lang = input("\nEnter 'fr'- French, 'es' - Spanish, 'ca' - Catalan, 'en' - English, 'fi' - Finnish: - ")
        userInput = input("\nPlease enter the name of the source doc you want translating ")
        userInput2 = input("\nPlease enter the name you want the ouput doc to be saved as ")

        p = paraText(userInput,lang)

        doc = Document()
        doc.add_paragraph(p)
        doc.save('{:}.docx'.format(userInput2))
        
        print("\nYour translated document can be found in the same folder as your source document")
    
    elif (userSelection == '3'):
        
        og, docs, first_doc = read_in_docx()

        y = input("What name would you like the destination file to have?: ")

        print("Here is the order of the documents: ", og)

        x = input("If this is correct please enter 'Y' to continue: ")

        if x == 'Y'.lower():
            combine_all_docx(first_doc,docs,y)
            print("The document has been created")
        else:
            print("Please contact the love and he will help you")
    
    else:
    
        #Handles any incorrect values entered into the menu and returns the user to the menu
        print("\nInvalid choice entered please enter a valid choice from the menu") 
    
    #repeats the menu upon completion of the chosen function
    userSelection = menu()  
